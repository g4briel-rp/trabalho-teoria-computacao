from docx import Document
import copy
import os

class MT_nao_deterministica():
    def __init__(self, estado_atual, palavra, fita, posicao_da_fita):
        self.estado_atual = estado_atual
        self.palavra = palavra
        self.fita = fita
        self.posicao_da_fita = posicao_da_fita
        self.historico = []
    
    def getEstadoAtual(self):
        return self.estado_atual

    def getPalavra(self):
        return self.palavra
    
    def getFita(self):
        return ''.join(self.fita)
    
    def getPosicaoDaFita(self):
        return int(self.posicao_da_fita)
    
    def getHistorico(self):
        return self.historico
    
    def adicionar(self, transicao):
        self.historico.append(transicao)

    def le(self, posicao = None):
        if posicao is None:
            return self.fita[self.posicao_da_fita]
        else:
            return self.fita[posicao]

    def escreve(self, proxEstado, vai_escrever, acao):
        if self.posicao_da_fita < 0 or self.posicao_da_fita > len(self.palavra):
            print('Excedeu o limite da fita')
            exit()

        self.estado_atual = proxEstado

        self.fita[self.posicao_da_fita] = vai_escrever

        if acao == 'R':
            self.posicao_da_fita += 1
        elif acao == 'L':
            self.posicao_da_fita -= 1
        
        print('Palavra: ', self.getPalavra())
        print('Fita: ', self.getFita())
        print('Nova posição da fita: ', self.getPosicaoDaFita())

def MT_deterministica():
    # -------------------------------------- LEITURA DO ARQUIVO E INICIAÇÃO DAS VARIÁVEIS -----------------------------------------
    doc = Document('MT-deterministica.docx')

    # para acessar individualmente uma posição do array, usar array[num]
     
    # estados possíveis da MT
    estados = doc.paragraphs[0].text.split(',')

    alfabeto_de_entrada = doc.paragraphs[1].text.split(',')

    alfabeto_da_fita = doc.paragraphs[2].text.split(',')

    # para acessar individualmente uma posição do array, usar array[num][num]
    #col1               col2        col3            col4        col5 
    #estadoAtual        lendo       proxEstado      escrevo     paraOndeVou (R ou L)

    # todas as transições da MT
    transicoes = [paragrafo.text.replace(";", ",").split(',') for paragrafo in doc.paragraphs[3:23]]

    # especificação dos estados da MT
    estado_inicial = doc.paragraphs[23].text.split(',')
    estado_aceitacao = doc.paragraphs[24].text.split(',')
    estado_rejeicao = doc.paragraphs[25].text.split(',')

    entrada = input("Digite o nome do arquivo de leitura(incluindo a extensão): ")
    doc2 = Document(entrada)
    # entradas da MT
    entradas = [paragrafo.text for paragrafo in doc2.paragraphs]
    # -----------------------------------------------------------------------------------------------------------------------------

    # ------------------------------------------------------ ALGORÍTMO ------------------------------------------------------------

    for palavras in entradas:

        flag = False # quando a palavra for aceita: flag = True
        estado_atual = estado_inicial[0]
        posicao_da_fita = 0
        auxLeitura = 0 # # auxLeitura só vai ser 1 quando já tiver feito alguma ação para esquerda 'L'
        copia = copy.copy(palavras)

        os.system('clear')

        # enquando n chegar no estado de aceitação
        while not flag:
            palavras = list(palavras)
            print("posicao atual da fita: ", posicao_da_fita,"tamanho: ", len(palavras))

            # para não deixar ler '#' direto. Transição de q0 para q5
            # auxLeitura só vai ser 1 quando já tiver feito alguma ação para esquerda 'L'
            if(auxLeitura == 0 and palavras[posicao_da_fita] not in alfabeto_de_entrada):
                os.system('clear')
                print("Leitura de um caracter inválido!\n")
                input("Pressione Enter para continuar...\n")
                break

            if palavras[posicao_da_fita] in alfabeto_da_fita: # se o caractere de leitura pertence ao alfabeto da fita
                existe_transicao = False

                for t in transicoes:
                    # se a transição é valida (todos os seus valores existem)
                    if t[0] == estado_atual and t[1] == palavras[posicao_da_fita] and t[2] in estados and t[3] in alfabeto_da_fita:
                        aux = t
                        existe_transicao = True
                        break
                
                if existe_transicao:
                    palavras[posicao_da_fita] = aux[3] # sobrescreve o caractere atual
                    if posicao_da_fita >= 0 and posicao_da_fita <= len(palavras): # se a posição da fita é válida -> realizar as ações
                        if aux[4] == 'R':
                            posicao_da_fita += 1
                        elif aux[4] == 'L':
                            auxLeitura = 1
                            posicao_da_fita -= 1 
                        print("estado atual: ", estado_atual," le: ",aux[1]," prox est: ",aux[2]," escrevo: ",aux[3], " acao: ", aux[4])
                        estado_atual = aux[2]
                    else:
                        print("\nFora do intervalo da fita permitido!\n")
                        input("Pressione Enter para continuar...\n")
                        exit()
                else: # caso a transição n seja válida
                    os.system('clear')
                    print("Não existe transição do estado ", estado_atual, " lendo ", palavras[posicao_da_fita])
                    print("Posição de parada da cabeça de leitura:", posicao_da_fita)
                    print("Palavra rejeitada pela máquina")
                    input("Pressione Enter para continuar...\n")
                    break
                
                # se está no estado de aceitação && for a ultima posicao da palavra && o ultimo caracter for '_' -> atualiza flag e aceita a palavra
                if estado_atual in estado_aceitacao:
                    if posicao_da_fita == len(palavras):
                        if palavras[posicao_da_fita - 1] == '_':
                            flag = True
                            os.system('clear')
                            print("A palavra foi aceita pela máquina")
                 
                print("palavra original: ", copia, "\napos alteração: ", ''.join(palavras))
                input("Pressione Enter para continuar...\n")
            else:
                print("Entrada com caracter não pertencente ao alfabeto!")
                input("Pressione Enter para continuar...\n")
                break

def mt_nao_deterministica():

    # -------------------------------------- LEITURA DO ARQUIVO E INICIAÇÃO DAS VARIÁVEIS -----------------------------------------
    doc = Document(('MT-nao-deterministica.docx'))
    estados = doc.paragraphs[0].text.split(',')
    alfabeto_de_entrada = doc.paragraphs[1].text.split(',')
    alfabeto_da_fita = doc.paragraphs[2].text.split(',')

    # para acessar individualmente uma posição do array, usar array[num][num]
    #col1               col2        col3            col4        col5 
    #estadoAtual        lendo       proxEstado      escrevo     paraOndeVou (R ou L)
    transicoes = [paragrafo.text.replace(";", ",").split(',') for paragrafo in doc.paragraphs[3:30]]

    estado_inicial = doc.paragraphs[30].text.split(',')
    estado_aceitacao = doc.paragraphs[31].text.split(',')
    estado_rejeicao = doc.paragraphs[32].text.split(',')

    entrada = input("Digite o nome do arquivo de leitura(incluindo a extensão): ")
    doc2 = Document(entrada)
    
    # entradas da MT
    entradas = [paragrafo.text for paragrafo in doc2.paragraphs]

    # -----------------------------------------------------------------------------------------------------------------------------

    # ------------------------------------------------------ ALGORÍTMO ------------------------------------------------------------

    for palavra in entradas:
        # Quando chega no estado de aceitação: flag = True
        flag = False
        posicao_da_fita_visitada = []
        maquina = MT_nao_deterministica(estado_inicial[0], palavra, list(palavra), 0)

        os.system('clear')

        # Enquanto não chegar no estado de aceitação
        while not flag:
            print("posicao atual da fita: ", maquina.getPosicaoDaFita())
            transicoes_encontradas = []

            # para saber se a posição ja foi visitada
            # para saber qual alfabeto vai utilizar
            if maquina.getPosicaoDaFita() not in posicao_da_fita_visitada:
                posicao_da_fita_visitada.append(maquina.getPosicaoDaFita())

                for t in transicoes:
                    if t[2] in estados: # se o próximo estado existe
                        if t[2] in estado_rejeicao: # se o próximo estado pertence ao estado de rejeição
                            if maquina.getEstadoAtual() == t[0]: #  se o estado atual da máquina existe
                                # se o que se está lendo existe e está no alfabeto de entrada
                                if maquina.le() == t[1] and maquina.le() in alfabeto_de_entrada:
                                    if t[3] in alfabeto_da_fita: # se o que se está escrevende existe no alfabeto da fita
                                        # se tudo for válido, guardar a transição
                                        transicoes_encontradas.append(t)
                        elif t[2] in estado_aceitacao:# se o próximo estado pertence ao estado de aceitação
                            if maquina.getEstadoAtual() == t[0]:
                                if maquina.le() == t[1] and maquina.le() in alfabeto_da_fita:
                                    if t[3] in alfabeto_da_fita:
                                        transicoes_encontradas.append(t)
                
                if len(transicoes_encontradas) > 0:
                    print(len(transicoes_encontradas), ' transição(ões) encontrada(s)')
                    if len(transicoes_encontradas) == 1:# caso o estado tenha uma transição para o caractere que esteja lendo
                        # a função escreve recebe: pŕoximo estado, o que vai escrever, ação
                        # de acordo com isso atualiza a fita
                        maquina.escreve(transicoes_encontradas[0][2], transicoes_encontradas[0][3], transicoes_encontradas[0][4])
                    else:
                        # caso o estado tenha mais de uma transição para o caractere que esteja lendo
                        # sabemos a linguagem da máquina (WW)
                        # a ideia é utilizar uma transição até que chegue na metade da fita
                        # só quando chegar na metade, marcar o caractere
                        if maquina.getPosicaoDaFita() < int(len(palavra) / 2):
                            # transicoes_encontradas:
                            # [['10', 'a', '10', 'a', 'R'], ['10', 'a', '2', 'Y', 'L']] -> para o estado q10
                            # [['9', 'b', '2', 'Y', 'L'], ['9', 'b', '9', 'b', 'R']] -> para o estado q9

                            # se tiver antes da metade, usa a primeira posição
                            if maquina.le() == 'a':
                                maquina.escreve(transicoes_encontradas[0][2], transicoes_encontradas[0][3], transicoes_encontradas[0][4])
                            else:
                                maquina.escreve(transicoes_encontradas[1][2], transicoes_encontradas[1][3], transicoes_encontradas[1][4])
                        else:
                            # transicoes_encontradas:
                            # [['10', 'a', '10', 'a', 'R'], ['10', 'a', '2', 'Y', 'L']] -> para o estado q10
                            # [['9', 'b', '2', 'Y', 'L'], ['9', 'b', '9', 'b', 'R']] -> para o estado q9

                            # se tiver depois da metade, usa a segunda posição
                            if maquina.le() == 'a':
                                maquina.escreve(transicoes_encontradas[1][2], transicoes_encontradas[1][3], transicoes_encontradas[1][4])
                            else:
                                maquina.escreve(transicoes_encontradas[0][2], transicoes_encontradas[0][3], transicoes_encontradas[0][4])

                    input("Pressione Enter para continuar...\n")

                    # se esta no estado de aceitação, se a posição da fita for a ultima e estiver lendo '_'
                    # atualiza a flag para para o loop e aceita a palavra
                    if maquina.getEstadoAtual() in estado_aceitacao:
                        if maquina.getPosicaoDaFita() == len(maquina.getPalavra()):
                            if maquina.le(maquina.getPosicaoDaFita()-1) == '_':
                                flag = True
                                print("A palavra foi aceita pela máquina")
                                input("Pressione Enter para continuar...\n")
                else: # caso não tenha encontrado nenhuma transição
                    print('Nenhuma transição encontrada')
                    print('Palavra rejeitada pela máquina')
                    input("Pressione Enter para continuar...\n")
                    break
            else:
                # quando as posições da fita ja foram visitadas, pode-se usar a ação 'L'
                # só entra aqui quando está em uma posição da fita que ja foi visitada
                # o restante é igual ao do if anterior
                for t in transicoes:
                    if t[2] in estados:
                        if maquina.getEstadoAtual() == t[0]:
                            if maquina.le() == t[1] and maquina.le() in alfabeto_da_fita:
                                if t[3] in alfabeto_da_fita:
                                    transicoes_encontradas.append(t)
                
                if len(transicoes_encontradas) > 0:
                    print(len(transicoes_encontradas), ' transição(ões) encontrada(s)')
                    if len(transicoes_encontradas) == 1:
                        maquina.escreve(transicoes_encontradas[0][2], transicoes_encontradas[0][3], transicoes_encontradas[0][4])
                    else:
                        if maquina.getPosicaoDaFita() < int(len(palavra) / 2):
                            if maquina.le() == 'a':
                                maquina.escreve(transicoes_encontradas[0][2], transicoes_encontradas[0][3], transicoes_encontradas[0][4])
                            else:
                                maquina.escreve(transicoes_encontradas[1][2], transicoes_encontradas[1][3], transicoes_encontradas[1][4])
                        else:
                            if maquina.le() == 'a':
                                maquina.escreve(transicoes_encontradas[1][2], transicoes_encontradas[1][3], transicoes_encontradas[1][4])
                            else:
                                maquina.escreve(transicoes_encontradas[0][2], transicoes_encontradas[0][3], transicoes_encontradas[0][4])

                    input("Pressione Enter para continuar...\n")

                    if maquina.getEstadoAtual() in estado_aceitacao:
                        if maquina.getPosicaoDaFita() == len(maquina.getPalavra()):
                            if maquina.le(maquina.getPosicaoDaFita()-1) == '_':
                                flag = True
                                print("A palavra foi aceita pela máquina")
                                input("Pressione Enter para continuar...\n")
                else:
                    print('Nenhuma transição encontrada')
                    print('Palavra rejeitada pela máquina')
                    input("Pressione Enter para continuar...\n")
                    break

if __name__ == '__main__':
    escolha = input("Informe a maquina que se deseja utilizar:\n1- MT Determinística\n2- MT Não Determinística\n>>")
    print(escolha)

    # match == switch
    match escolha:
        case '1':
            MT_deterministica()
        case '2':
            mt_nao_deterministica()
        case _:
            print("Opção inválida!\n")

    