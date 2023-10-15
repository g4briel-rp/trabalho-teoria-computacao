from docx import Document
import copy
import os

class MT_nao_deterministica():
    def __init__(self, estado_atual, palavra, fita, posicao_da_fita):
        self.estado_atual = estado_atual
        self.palavra = palavra
        self.fita = fita
        self.posicao_da_fita = posicao_da_fita
        self.historico = []
    
    def getEstadoAtual(self):
        return self.estado_atual

    def getPalavra(self):
        return self.palavra
    
    def getFita(self):
        return ''.join(self.fita)
    
    def getPosicaoDaFita(self):
        return int(self.posicao_da_fita)
    
    def getHistorico(self):
        return self.historico
    
    def adicionar(self, transicao):
        self.historico.append(transicao)

    def le(self, posicao = None):
        if posicao is None:
            return self.fita[self.posicao_da_fita]
        else:
            return self.fita[posicao]

    def escreve(self, proxEstado, vai_escrever, acao):
        if self.posicao_da_fita < 0 or self.posicao_da_fita > len(self.palavra):
            print('Excedeu o limite da fita')
            exit()

        self.estado_atual = proxEstado

        self.fita[self.posicao_da_fita] = vai_escrever

        if acao == 'R':
            self.posicao_da_fita += 1
        elif acao == 'L':
            self.posicao_da_fita -= 1
        
        print('Palavra: ', self.getPalavra())
        print('Fita: ', self.getFita())
        print('Nova posição da fita: ', self.getPosicaoDaFita())

def mt_nao_deterministica():
    doc = Document(('MT-nao-deterministica.docx'))
    estados = doc.paragraphs[0].text.split(',')
    alfabeto_de_entrada = doc.paragraphs[1].text.split(',')
    alfabeto_da_fita = doc.paragraphs[2].text.split(',')

    #array[num][num]
    #col1-col2-col3-col4-col5 
    #estadoAtual-lendo-proxEstado-escrevo-acao
    transicoes = [paragrafo.text.replace(";", ",").split(',') for paragrafo in doc.paragraphs[3:30]]

    estado_inicial = doc.paragraphs[30].text.split(',')
    estado_aceitacao = doc.paragraphs[31].text.split(',')
    estado_rejeicao = doc.paragraphs[32].text.split(',')

    entrada = input("Digite o nome do arquivo de leitura(incluindo a extensão): ")
    doc2 = Document(entrada)
    
    # entradas da MT
    entradas = [paragrafo.text for paragrafo in doc2.paragraphs]

    for palavra in entradas:
        flag = False
        posicao_da_fita_visitada = []
        maquina = MT_nao_deterministica(estado_inicial[0], palavra, list(palavra), 0)

        os.system('clear')

        while not flag:
            print("posicao atual da fita: ", maquina.getPosicaoDaFita())
            transicoes_encontradas = []

            if maquina.getPosicaoDaFita() not in posicao_da_fita_visitada:
                posicao_da_fita_visitada.append(maquina.getPosicaoDaFita())

                for t in transicoes:
                    if t[2] in estados:
                        if t[2] in estado_rejeicao:
                            if maquina.getEstadoAtual() == t[0]:
                                if maquina.le() == t[1] and maquina.le() in alfabeto_de_entrada:
                                    if t[3] in alfabeto_da_fita:
                                        transicoes_encontradas.append(t)
                        elif t[2] in estado_aceitacao:
                            if maquina.getEstadoAtual() == t[0]:
                                if maquina.le() == t[1] and maquina.le() in alfabeto_da_fita:
                                    if t[3] in alfabeto_da_fita:
                                        transicoes_encontradas.append(t)
                
                if len(transicoes_encontradas) > 0:
                    print(len(transicoes_encontradas), ' transição(ões) encontrada(s)')
                    if len(transicoes_encontradas) == 1:
                        maquina.escreve(transicoes_encontradas[0][2], transicoes_encontradas[0][3], transicoes_encontradas[0][4])
                    else:
                        if maquina.getPosicaoDaFita() < int(len(palavra) / 2):
                            if maquina.le() == 'a':

                                maquina.escreve(transicoes_encontradas[0][2], transicoes_encontradas[0][3], transicoes_encontradas[0][4])
                            else:
                                maquina.escreve(transicoes_encontradas[1][2], transicoes_encontradas[1][3], transicoes_encontradas[1][4])
                        else:
                            if maquina.le() == 'a':
                                maquina.escreve(transicoes_encontradas[1][2], transicoes_encontradas[1][3], transicoes_encontradas[1][4])
                            else:
                                maquina.escreve(transicoes_encontradas[0][2], transicoes_encontradas[0][3], transicoes_encontradas[0][4])

                    input("Pressione Enter para continuar...\n")

                    if maquina.getEstadoAtual() in estado_aceitacao:
                        if maquina.getPosicaoDaFita() == len(maquina.getPalavra()):
                            if maquina.le(maquina.getPosicaoDaFita()-1) == '_':
                                flag = True
                                print("A palavra foi aceita pela máquina")
                                input("Pressione Enter para continuar...\n")
                else:
                    print('Nenhuma transição encontrada')
                    print('Palavra rejeitada pela máquina')
                    input("Pressione Enter para continuar...\n")
                    break
            else:
                for t in transicoes:
                    if t[2] in estados:
                        if maquina.getEstadoAtual() == t[0]:
                            if maquina.le() == t[1] and maquina.le() in alfabeto_da_fita:
                                if t[3] in alfabeto_da_fita:
                                    transicoes_encontradas.append(t)
                
                if len(transicoes_encontradas) > 0:
                    print(len(transicoes_encontradas), ' transição(ões) encontrada(s)')
                    if len(transicoes_encontradas) == 1:
                        maquina.escreve(transicoes_encontradas[0][2], transicoes_encontradas[0][3], transicoes_encontradas[0][4])
                    else:
                        if maquina.getPosicaoDaFita() < int(len(palavra) / 2):
                            if maquina.le() == 'a':
                                maquina.escreve(transicoes_encontradas[0][2], transicoes_encontradas[0][3], transicoes_encontradas[0][4])
                            else:
                                maquina.escreve(transicoes_encontradas[1][2], transicoes_encontradas[1][3], transicoes_encontradas[1][4])
                        else:
                            if maquina.le() == 'a':
                                maquina.escreve(transicoes_encontradas[1][2], transicoes_encontradas[1][3], transicoes_encontradas[1][4])
                            else:
                                maquina.escreve(transicoes_encontradas[0][2], transicoes_encontradas[0][3], transicoes_encontradas[0][4])

                    input("Pressione Enter para continuar...\n")

                    if maquina.getEstadoAtual() in estado_aceitacao:
                        if maquina.getPosicaoDaFita() == len(maquina.getPalavra()):
                            if maquina.le(maquina.getPosicaoDaFita()-1) == '_':
                                flag = True
                                print("A palavra foi aceita pela máquina")
                                input("Pressione Enter para continuar...\n")
                else:
                    print('Nenhuma transição encontrada')
                    print('Palavra rejeitada pela máquina')
                    input("Pressione Enter para continuar...\n")
                    break

if __name__ == '__main__':
    mt_nao_deterministica()